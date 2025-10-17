# -*- coding: utf-8 -*-
"""
Doc Evaluator (Final Refactored)
- e5-large 임베딩 기반 문장 유사도(코사인) 반영
- ngram + cosine redundancy 혼합
- Relevance / Consistency / Weighted Final Score 통합
- Markdown + JSON + CSV 출력
"""
import os, re, json, math, unicodedata
from typing import List, Dict, Any
from collections import namedtuple, Counter, defaultdict

# --------------------------
# 의존성
# --------------------------
try:
    import docx
    import numpy as np
    from sklearn.metrics.pairwise import cosine_similarity
    from sentence_transformers import SentenceTransformer
    # BM25를 쓰지 않으므로 import 생략 가능. 남겨두어도 무방.
    # from rank_bm25 import BM25Okapi
except ImportError as e:
    raise RuntimeError(
        "필수 모듈이 설치되지 않았습니다. pip install python-docx sentence-transformers scikit-learn"
    ) from e

# --------------------------
# 기본 유틸
# --------------------------
def _norm(s: str) -> str:
    s = unicodedata.normalize("NFKC", (s or "").strip())
    s = re.sub(r"\s+", " ", s)
    return s

def split_ko_sentences(text: str) -> List[str]:
    text = re.sub(r"\s+", " ", text or "")
    # 한국어 문장 경계 간단 탐지(마침표/물음표/느낌표/“다.” 패턴 포함)
    sents = re.split(r"(?<=[\.?!])\s+|(?<=다\.)\s+", text)
    return [s.strip() for s in sents if s.strip()]

_TOKEN_RE = re.compile(r"[A-Za-z가-힣0-9%\.]+", re.UNICODE)
def _simple_tokenize(text: str):
    if not text:
        return []
    return [t for t in _TOKEN_RE.findall(text.lower()) if t]

# --------------------------
# 핵심 지표 계산 함수
# --------------------------
def keyword_overlap(a: str, b: str) -> float:
    ta = set(re.findall(r"[가-힣A-Za-z0-9]+", (a or "").lower()))
    tb = set(re.findall(r"[가-힣A-Za-z0-9]+", (b or "").lower()))
    if not ta or not tb:
        return 0.0
    return len(ta & tb) / len(ta | tb)

def ngram_redundancy(sentences: List[str], n: int = 3) -> float:
    grams = []
    for s in sentences:
        toks = re.findall(r"[가-힣A-Za-z0-9]+", (s or "").lower())
        if len(toks) >= n:
            grams += list(zip(*[toks[i:] for i in range(n)]))
    if not grams:
        return 0.0
    c = Counter(grams)
    dup = sum(v - 1 for v in c.values() if v > 1)
    return dup / (len(grams) + 1e-6)

# --------- 전역 임베더 캐시(문장 유사도 계산 속도 개선) ----------
_EMBEDDER_CACHE = {"name": None, "model": None}

def _get_embedder(model_name: str = "intfloat/e5-large") -> SentenceTransformer:
    global _EMBEDDER_CACHE
    if _EMBEDDER_CACHE["model"] is not None and _EMBEDDER_CACHE["name"] == model_name:
        return _EMBEDDER_CACHE["model"]
    # 모델 최초/재로딩
    model = SentenceTransformer(model_name)
    _EMBEDDER_CACHE["name"] = model_name
    _EMBEDDER_CACHE["model"] = model
    return model

def cosine_redundancy(sentences: List[str], model_name: str = "intfloat/e5-large", threshold: float = 0.9) -> float:
    """
    코사인 유사도 기반 반복률 계산:
    - 문장쌍 코사인 유사도 >= threshold(기본 0.9)인 비율
    """
    sentences = [s for s in sentences if s.strip()]
    if len(sentences) < 2:
        return 0.0
    try:
        model = _get_embedder(model_name)
        embeddings = model.encode(sentences, normalize_embeddings=True)
        sims = cosine_similarity(embeddings)
    except Exception as e:
        print(f"[cosine_redundancy] 임베딩 실패: {e}")
        return 0.0

    n = len(sentences)
    total_pairs, high_pairs = 0, 0
    # 상삼각만 카운트(i<j)
    for i in range(n):
        for j in range(i + 1, n):
            total_pairs += 1
            if sims[i, j] >= threshold:
                high_pairs += 1
    return high_pairs / total_pairs if total_pairs else 0.0

def simple_coherence(sentences: List[str]) -> float:
    if len(sentences) < 2:
        return 0.5
    scores = [keyword_overlap(sentences[i], sentences[i + 1]) for i in range(len(sentences) - 1)]
    return sum(scores) / len(scores)

def _fluency(sents: List[str]) -> float:
    if not sents:
        return 0.5
    lens = [len(s) for s in sents]
    mean_len = sum(lens) / len(lens)
    punct = sum(ch in ".,;:?!~" for s in sents for ch in s) / (sum(lens) + 1e-6)
    score = 0.5 + 0.5 * math.tanh((mean_len - 25) / 50) - 0.2 * abs(punct - 0.03)
    return max(0.0, min(1.0, score))

def relevance_score(section_text: str, required_title: str) -> float:
    return keyword_overlap(section_text, required_title)

def consistency_score(section_text: str) -> float:
    # 단위 일관성: 등장한 단위 타입의 다양도가 낮을수록(=일관성 높음) 가점
    nums = re.findall(r"\d+(?:[\.,]\d+)?\s?(%|ms|초|일|주|개월|월|분기|년|원|만원|억)?", section_text or "")
    units = [u.strip() for u in nums if u and u.strip()]
    return len(set(units)) / len(units) if units else 1.0

# --------------------------
# 문서 파서
# --------------------------
class DocParser:
    def parse(self, docx_path: str):
        if not os.path.exists(docx_path):
            return None
        try:
            d = docx.Document(docx_path)
        except Exception:
            return None
        paras = [p.text.strip() for p in d.paragraphs if p.text and p.text.strip()]
        sections, cur_title, cur_buf = [], None, []
        for p in d.paragraphs:
            style = getattr(p.style, "name", "") or ""
            text = (p.text or "").strip()
            if not text:
                continue
            # Heading / "제목" 스타일이면 새 섹션 시작
            if style.startswith("Heading") or "제목" in style:
                if cur_title is not None or cur_buf:
                    sections.append({"title": _norm(cur_title), "text": _norm("\n".join(cur_buf))})
                cur_title, cur_buf = text, []
            else:
                cur_buf.append(text)
        if cur_title is not None or cur_buf:
            sections.append({"title": _norm(cur_title), "text": _norm("\n".join(cur_buf))})
        full_text = _norm("\n".join(paras))
        sentences = split_ko_sentences(full_text)
        Doc = namedtuple("Doc", ["sections", "paragraphs", "sentences", "text"])
        return Doc(sections=sections, paragraphs=paras, sentences=sentences, text=full_text)

# --------------------------
# 섹션 평가
# --------------------------
def evaluate_section(rt_title: str, sec_text: str) -> Dict[str, Any]:
    sents = split_ko_sentences(sec_text)
    if not (sec_text or "").strip():
        return {"required_title": rt_title, "exists": False, "final": 0.0}

    # 간단 정확도 대용치(길이 기반)
    accuracy = 0.8 if len(sec_text) > 200 else 0.4
    flu = _fluency(sents)
    coh = simple_coherence(sents)

    # 반복률: ngram + cosine(0.9) 합성
    red_ngram = ngram_redundancy(sents, n=3)
    red_cosine = cosine_redundancy(sents, model_name="intfloat/e5-large", threshold=0.9)
    redundancy = 0.5 * red_ngram + 0.5 * red_cosine

    # 추가 지표
    relevance = relevance_score(sec_text, rt_title)
    consistency = consistency_score(sec_text)

    # 최종 가중합
    final = (
        0.25 * accuracy
        + 0.20 * relevance
        + 0.20 * coh
        + 0.15 * flu
        + 0.10 * consistency
        + 0.10 * (1 - redundancy)
    )

    return {
        "required_title": rt_title,
        "exists": True,
        "accuracy": accuracy,
        "fluency": flu,
        "coherence": coh,
        "redundancy": redundancy,
        "relevance": relevance,
        "consistency": consistency,
        "final": final,
    }

# --------------------------
# 보고서 생성
# --------------------------
def run_combined_report(docx_paths: List[str], required_titles: List[str], cfg: Dict[str, Any], out_path: str):
    parser = DocParser()
    results = defaultdict(list)

    for path in docx_paths:
        doc = parser.parse(path)
        if not doc:
            continue
        for rt in required_titles:
            # 섹션 제목이 rt를 포함하는 첫 섹션
            sec = next((s for s in doc.sections if s.get("title") and rt in s["title"]), None)
            text = sec["text"] if sec else ""
            res = evaluate_section(rt, text)
            res["doc"] = os.path.basename(path)
            results[rt].append(res)

    # 평균 요약(누락 섹션 제외)
    summary = {}
    for rt, vals in results.items():
        valid = [v for v in vals if v.get("exists")]
        if not valid:
            continue
        def avg(k): return sum(v[k] for v in valid) / len(valid)
        summary[rt] = {
            "accuracy": avg("accuracy"),
            "relevance": avg("relevance"),
            "coherence": avg("coherence"),
            "fluency": avg("fluency"),
            "consistency": avg("consistency"),
            "redundancy": avg("redundancy"),
            "final": avg("final"),
        }

    # Markdown 보고서
    lines = []
    lines.append("# 통합 문서 평가 보고서\n\n")
    lines.append("| 제목 | Acc | Rel | Coh | Flu | Cons | Red(↓) | Final |\n")
    lines.append("|---|---:|---:|---:|---:|---:|---:|---:|\n")
    for rt in required_titles:
        s = summary.get(rt)
        if not s:
            lines.append(f"| {rt} | 0 | 0 | 0 | 0 | 0 | 0 | 0 |\n")
        else:
            lines.append(
                f"| {rt} | {s['accuracy']:.2f} | {s['relevance']:.2f} | {s['coherence']:.2f} "
                f"| {s['fluency']:.2f} | {s['consistency']:.2f} | {s['redundancy']:.2f} | {s['final']:.2f} |\n"
            )

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("".join(lines))

    # JSON, CSV 저장
    base, _ = os.path.splitext(out_path)
    json_path = base + "_summary.json"
    csv_path = base + "_summary.csv"
    with open(json_path, "w", encoding="utf-8") as jf:
        json.dump(summary, jf, ensure_ascii=False, indent=2)
    try:
        import pandas as pd
        df = pd.DataFrame([{**{"section": k}, **v} for k, v in summary.items()])
        df.to_csv(csv_path, index=False, encoding="utf-8-sig")
    except Exception:
        pass

    print(" Markdown:", out_path)
    print(" JSON:", json_path)
    print(" CSV :", csv_path)
    return out_path, json_path, csv_path

# --------------------------
# 실행 예시 (Colab/로컬 공통)
# --------------------------
REQUIRED_TITLES = [
    "연구개발 목표","연구개발 내용","연구개발성과 활용계획 및 기대효과","연구기획과제의 개요",
    "연구개발과제의 배경","연구개발과제의 필요성","보안등급의 분류 및 해당 사유","기술개발 핵심어(키워드)",
    "연차별 개발목표","연차별 개발내용 및 범위","추진방법 및 전략","과제 성과의 활용방안",
    "신규사업 신설의 기대효과","사회적 가치 창출 계획","사회적 가치창출의 기대효과",
    "경제적 성과창출의 기대효과","신규 인력 채용 계획 및 활용 방안"
]

CONFIG = {
    "models": {
        "embed": "BAAI/bge-m3",   # (현재 코드는 redundancy에 e5-large 사용 중, 필요시 바꿀 수 있음)
        "nli": "rule-lite",
        "qna": "rule-lite"
    }
}

if __name__ == "__main__":
    # Colab 예시 경로 (Google Drive 마운트 가정)
    DOCX_LIST = [ 
        "/content/drive/MyDrive/e5L/section_1_연구기획과제의 개요.docx",
        "/content/drive/MyDrive/e5L/section_2_연구개발 내용.docx",
        "/content/drive/MyDrive/e5L/section_3_연구개발성과 활용계획 및 기대효과.docx",
        "/content/drive/MyDrive/e5L/section_4_연구기획과제의 개요.docx",
        "/content/drive/MyDrive/e5L/section_5_연구개발과제의 배경.docx",
        "/content/drive/MyDrive/e5L/section_6_연구개발과제의 필요성.docx",
        "/content/drive/MyDrive/e5L/section_7_보안등급의 분류 및 해당 사유.docx",
        "/content/drive/MyDrive/e5L/section_8_기술개발 핵심어(키워드).docx",
        "/content/drive/MyDrive/e5L/section_9_연차별 개발목표.docx",
        "/content/drive/MyDrive/e5L/section_10_연차별 개발내용 및 범위.docx",
        "/content/drive/MyDrive/e5L/section_11_추진방법 및 전략.docx",
        "/content/drive/MyDrive/e5L/section_12_과제 성과의 활용방안.docx",
        "/content/drive/MyDrive/e5L/section_13_신규사업 신설의 기대효과.docx",
        "/content/drive/MyDrive/e5L/section_14_사회적 가치 창출 계획.docx",
        "/content/drive/MyDrive/e5L/section_15_사회적 가치창출의 기대효과.docx",
        "/content/drive/MyDrive/e5L/section_16_경제적 성과창출의 기대효과.docx",
        "/content/drive/MyDrive/e5L/section_17_신규 인력 채용 계획 및 활용 방안.docx"
    ]

    if not DOCX_LIST:
        print("[알림] DOCX_LIST에 문서 경로를 17개 넣어 실행하세요.")
    else:
        out_path = "/content/drive/MyDrive/e5L/e5v1.md"
        os.makedirs(os.path.dirname(out_path), exist_ok=True)
        md_path, json_path, csv_path = run_combined_report(
            DOCX_LIST, REQUIRED_TITLES, CONFIG, out_path=out_path
        )
        print(" 보고서 저장 완료:", md_path)
        print(" 섹션 요약 JSON:", json_path)
        print(" 섹션 요약 CSV :", csv_path)
