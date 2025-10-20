# -*- coding: utf-8 -*-
# LangGraph 기반 구조화 버전 (로컬 오프라인용)
# =====================================================================
from langgraph.graph import StateGraph, END
from transformers import AutoTokenizer, AutoModelForCausalLM, pipeline
from docx import Document
import json, re, torch

# ===== 1. 환경 변수 =====
GUIDELINE_FILE = "/home/alpaco/autosry/rnd_guideline.json" 
RAG_JSON_FILES = ["/home/alpaco/autosry/rag_chunks500_50.json"] 
GEN_NAME = "skt/A.X-4.0-Light"
DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
GEN_MAX_NEW_TOKENS = 5000
GEN_DO_SAMPLE = False
GEN_TEMPERATURE = None

# ===== 2. LLM 로드 =====
tok = AutoTokenizer.from_pretrained(GEN_NAME, trust_remote_code=True)
model = AutoModelForCausalLM.from_pretrained(GEN_NAME, trust_remote_code=True, device_map="auto").eval()
pipe = pipeline("text-generation", model=model, tokenizer=tok)

# ===== 3. 유틸리티 함수 =====
def first_n_lines(text: str, n_chars=350):
    t = " ".join(str(text).split())
    return t[:n_chars]

def clean_generated_text(text: str) -> str:
    """
    불필요한 Markdown/특수문자(#, *, -, • 등) 및 중복 공백 제거.
    """
    # 1) 프롬프트 구분자 이후만 사용
    if "#=========== 출력" in text:
        text = text.split("#=========== 출력", 1)[-1]

    # 2) Markdown 및 리스트 기호 제거 (줄바꿈은 보존)
    text = re.sub(r'[•●▪▶◇◆□▪️▫️–]', ' ', text)
    text = re.sub(r'^\s*[-#*]+\s*', '', text, flags=re.MULTILINE)

    # 3) 여러 줄바꿈 정리 (2줄 이상 → 2줄로 유지)
    text = re.sub(r'\n{3,}', '\n\n', text)

    # 4) 문장부호 앞뒤 공백 정리
    text = re.sub(r'\s+([\.,;:])', r'\1', text)

    # 5) 불필요한 양쪽 공백 정리
    text = text.strip()
    return text

# ===== 4. Node 정의 =====
def form_search_node(state):
    """문서유형 기반 템플릿 검색"""
    query = state["query"]
    if "R&D" in query:
        state["template_path"] = "/templates/rnd_plan_v3.docx"
        state["law_meta"] = {"name": "행정업무의 운영 및 혁신에 관한 규정", "ver": "제34518호(2024.05.21)"}
    else:
        state["template_path"] = "/templates/default_form.docx"
    return state

def context_search_node(state):
    """RAG 문서 로드"""
    with open(RAG_FILE, "r", encoding="utf-8") as f:
        state["context"] = json.load(f)
    return state

def context_builder_node(state):
    """요약 및 핵심 포인트 추출"""
    ctx = state["context"]
    first_chunk = ctx["chunks"][0]["text"] if "chunks" in ctx else str(ctx)[:500]
    state["context_summary"] = f"주요 내용 요약: {first_chunk[:400]}"
    return state

def draft_writer_node(state):
    """문서 초안 작성"""
    prompt = f"""
    역할: 연구기획전문가
    문서 템플릿: {state['template_path']}
    법령: {state['law_meta']}
    요약: {state['context_summary']}
    과제명: {state['query']}
    #=========== 출력
    """
    out = pipe(prompt, max_new_tokens=1500)[0]["generated_text"]
    state["draft_text"] = clean_text(out)
    return state

def validator_node(state):
    """간단한 형식/정책 검증"""
    draft = state["draft_text"]
    issues = []
    if "성과" not in draft:
        issues.append("성과 항목 누락")
    if len(draft) < 200:
        issues.append("내용 부족")
    state["validation_pass"] = len(issues) == 0
    state["validation_report"] = {"issues": issues, "length": len(draft)}
    return state

def repairer_node(state):
    """검증 실패 시 자기수정"""
    if not state["validation_pass"]:
        feedback = ", ".join(state["validation_report"]["issues"])
        revised_prompt = f"다음 문제를 수정하세요: {feedback}\n\n{state['draft_text']}"
        revised = pipe(revised_prompt, max_new_tokens=1200)[0]["generated_text"]
        state["draft_text"] = clean_text(revised)
    return state

def exporter_node(state):
    """DOCX 파일 출력"""
    doc = Document()
    doc.add_heading("AI 행정·R&D 문서 자동생성기 결과", 0)
    doc.add_paragraph(state["draft_text"])
    out_path = "RND_Result.docx"
    doc.save(out_path)
    state["output_path"] = out_path
    print(f"[DONE] 문서 생성 완료: {out_path}")
    return state

# ===== 5. LangGraph 구성 =====
graph = StateGraph()
graph.add_node("form_search", form_search_node)
graph.add_node("context_search", context_search_node)
graph.add_node("context_builder", context_builder_node)
graph.add_node("draft_writer", draft_writer_node)
graph.add_node("validator", validator_node)
graph.add_node("repairer", repairer_node)
graph.add_node("exporter", exporter_node)

# Edge 정의
graph.add_edge("form_search", "context_search")
graph.add_edge("context_search", "context_builder")
graph.add_edge("context_builder", "draft_writer")
graph.add_edge("draft_writer", "validator")
graph.add_conditional_edges(
    "validator",
    lambda s: "exporter" if s["validation_pass"] else "repairer"
)
graph.add_edge("repairer", "draft_writer")
graph.add_edge("exporter", END)

# ===== 6. 실행 예시 =====
app = graph.compile()

if __name__ == "__main__":
    state = app.invoke({"query": "2025년도 R&D 계획서 작성"})
    print(state["tester_report"])
